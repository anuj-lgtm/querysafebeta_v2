"""
Training pipeline: extracts text from uploaded documents, chunks it,
generates embeddings, and builds a FAISS index for chatbot retrieval.

Optimised flow
──────────────
1. Classify files → text-based (PDF, DOCX, TXT) vs image-based (JPG, PNG …)
2. Extract text directly from text-based files (PyMuPDF / python-docx / open)
   • Scanned PDF pages (< 50 chars) fall back to Gemini vision.
3. Use Gemini vision only for images + scanned pages (concurrent calls).
4. Chunk all extracted text → embed → FAISS index.
"""

import os
import re
import json
import logging
import base64
import platform
import subprocess
import io
import time
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz
import faiss
import numpy as np
from PIL import Image
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from google import genai

from django.conf import settings
from django.utils import timezone
from user_querySafe.chatbot.embedding_model import get_embedding_model

logger = logging.getLogger(__name__)

# ── Paths ─────────────────────────────────────────────────────────────
BASE_DIR = os.path.join(settings.DATA_DIR, "documents")
PDF_DIR = os.path.join(BASE_DIR, "files_uploaded")
IMAGE_DIR = os.path.join(BASE_DIR, "files_images")
TEXT_DIR = os.path.join(BASE_DIR, "files_captions")
CHUNK_DIR = os.path.join(BASE_DIR, "files_chunks")
INDEX_DIR = os.path.join(BASE_DIR, "vector_index")
META_DIR = os.path.join(BASE_DIR, "chunk-metadata")

for folder in [PDF_DIR, IMAGE_DIR, TEXT_DIR, CHUNK_DIR, INDEX_DIR, META_DIR]:
    os.makedirs(folder, exist_ok=True)

# ── Gemini client ─────────────────────────────────────────────────────
PROJECT_ID = settings.PROJECT_ID
REGION = settings.REGION
client = genai.Client(vertexai=True, project=PROJECT_ID, location=REGION)

# ── File type definitions ─────────────────────────────────────────────
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp'}
TEXT_DOC_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}

# Minimum characters per PDF page to consider it text-based (not scanned)
MIN_TEXT_CHARS = 50

# ── Gemini vision prompt ──────────────────────────────────────────────
def _build_vision_prompt(image_data, mime_type="image/png"):
    return [{
        "role": "user",
        "parts": [
            {
                "text": (
                    "You are a visual analysis expert. Extract and describe every element "
                    "in the image including:\n"
                    "- Text (as-is)\n"
                    "- Tables (as plain readable text)\n"
                    "- Charts/graphs (with insights and data)\n"
                    "- Images/diagrams (detailed description)\n"
                    "Output must be clean, complete, and human-readable."
                )
            },
            {"inline_data": {"mime_type": mime_type, "data": image_data}},
        ]
    }]


# =====================================================================
# STEP 1 — Classify & extract text
# =====================================================================

def _extract_text_from_pdf(file_path):
    """Extract text from a PDF.  Returns (text, scanned_page_images).

    scanned_page_images is a list of (page_label, base64_png) for pages
    where extracted text was below MIN_TEXT_CHARS (likely scanned), or
    pages that have embedded images with limited text (charts/diagrams).
    """
    text_parts = []
    scanned_pages = []
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        page_text = page.get_text("text").strip()
        page_images = page.get_images(full=True)

        if len(page_text) >= MIN_TEXT_CHARS:
            text_parts.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
            # If page has embedded images and limited text, also send to
            # Gemini vision so charts/diagrams are described.
            if page_images and len(page_text) < 500:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                b64 = base64.b64encode(img_bytes).decode("utf-8")
                scanned_pages.append((f"page{page_num + 1}", b64))
        else:
            # Scanned / image-heavy page → render to image for Gemini
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            scanned_pages.append((f"page{page_num + 1}", b64))
    doc.close()
    return "\n".join(text_parts), scanned_pages


def _extract_text_from_docx(file_path):
    """Extract full text from a DOCX file."""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _extract_text_from_txt(file_path):
    """Read plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _image_to_base64(file_path):
    """Load an image file and return (base64_str, mime_type)."""
    ext = os.path.splitext(file_path)[1].lower()
    mime_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.gif': 'image/gif', '.bmp': 'image/bmp',
    }
    mime = mime_map.get(ext, 'image/png')
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return b64, mime


# =====================================================================
# STEP 2 — Gemini vision for images (concurrent)
# =====================================================================

def _caption_single_image(b64_data, mime_type, label):
    """Send one image to Gemini and return (label, caption_text)."""
    try:
        prompt = _build_vision_prompt(b64_data, mime_type)
        response = client.models.generate_content(
            model=settings.GEMINI_VISION_MODEL,
            contents=prompt,
        )
        return label, response.text.strip()
    except Exception as e:
        logger.warning("Gemini vision failed for %s: %s", label, e)
        return label, f"[Error extracting from {label}: {e}]"


def _caption_images_concurrent(image_items, max_workers=3):
    """Process a list of (label, b64, mime) tuples concurrently.
    Returns combined text string.
    """
    if not image_items:
        return ""
    results = {}
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {
            pool.submit(_caption_single_image, b64, mime, label): label
            for label, b64, mime in image_items
        }
        for future in as_completed(futures):
            label, caption = future.result()
            results[label] = caption
    # Reassemble in submission order
    ordered_labels = [item[0] for item in image_items]
    parts = []
    for label in ordered_labels:
        parts.append(f"\n--- {label} ---\n{results.get(label, '')}")
    return "\n".join(parts)


# =====================================================================
# STEP 3 — Chunking
# =====================================================================

def _chunk_text(raw_text, chunk_size=1500, chunk_overlap=200):
    """Split text into overlapping chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap,
    )
    return splitter.split_text(raw_text)


# =====================================================================
# STEP 4 — Embed & build FAISS index
# =====================================================================

def _embed_and_index(chatbot_id, chunk_records):
    """Generate embeddings and write FAISS index + metadata.

    chunk_records: list of {"content": str, "source": str}
    """
    if not chunk_records:
        logger.warning("No chunks to embed for chatbot %s", chatbot_id)
        return False

    texts = [r["content"] for r in chunk_records]

    print(f"  Generating embeddings for {len(texts)} chunks …")
    model = get_embedding_model()
    embeddings = model.encode(texts, show_progress_bar=False)
    dimension = embeddings.shape[1]

    index = faiss.IndexFlatL2(dimension)
    index.add(np.array(embeddings).astype("float32"))

    index_path = os.path.join(INDEX_DIR, f"{chatbot_id}-index.index")
    meta_path = os.path.join(META_DIR, f"{chatbot_id}-chunks.json")
    faiss.write_index(index, index_path)
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(chunk_records, f, indent=2)

    print(f"  ✓ FAISS index saved ({len(texts)} chunks, dim={dimension})")
    return True


# =====================================================================
# STEP 5 — Legacy DOC conversion (for .doc files only)
# =====================================================================

def _convert_doc_to_text(file_path):
    """Convert old-format .doc to text via LibreOffice/Win32COM, then read.
    Returns extracted text or None on failure.
    """
    system = platform.system().lower()
    temp_pdf = file_path + ".tmp.pdf"

    try:
        if system != "windows":
            abs_input = os.path.abspath(file_path)
            output_dir = os.path.dirname(os.path.abspath(temp_pdf))
            cmd = [
                'soffice', '--headless',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', output_dir,
                abs_input,
            ]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if proc.returncode == 0:
                generated = os.path.join(
                    output_dir,
                    os.path.splitext(os.path.basename(abs_input))[0] + '.pdf'
                )
                if os.path.exists(generated):
                    os.rename(generated, temp_pdf)
        else:
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                try:
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    doc.SaveAs(os.path.abspath(temp_pdf), FileFormat=17)
                    doc.Close()
                finally:
                    word.Quit()
                    pythoncom.CoUninitialize()
            except Exception as e:
                logger.warning("Win32COM .doc conversion failed: %s", e)
                return None

        # Now extract text from the converted PDF
        if os.path.exists(temp_pdf):
            text, scanned = _extract_text_from_pdf(temp_pdf)
            os.remove(temp_pdf)
            # For .doc scanned pages, we just return the text we have
            return text if text.strip() else None
        return None

    except Exception as e:
        logger.warning(".doc conversion failed for %s: %s", file_path, e)
        return None
    finally:
        if os.path.exists(temp_pdf):
            try:
                os.remove(temp_pdf)
            except OSError:
                pass


# =====================================================================
# MAIN PIPELINE
# =====================================================================

def process_pipeline(chatbot_id):
    """Full training pipeline for a chatbot."""
    print(f"\n🚀 Starting pipeline for chatbot: {chatbot_id}")
    start_time = time.time()

    # 1. Discover uploaded files ────────────────────────────────────────
    all_files = [
        f for f in os.listdir(PDF_DIR)
        if f.startswith(chatbot_id + "_")
    ]

    # Check if URL sources exist (don't bail out if only URLs)
    has_urls = False
    try:
        from user_querySafe.models import ChatbotURL
        has_urls = ChatbotURL.objects.filter(chatbot__chatbot_id=chatbot_id).exists()
    except Exception:
        pass

    if not all_files and not has_urls:
        print(f"  ❌ No files or URLs found for chatbot {chatbot_id}")
        from user_querySafe.models import Chatbot
        Chatbot.objects.filter(chatbot_id=chatbot_id).update(status="error")
        return

    print(f"  Found {len(all_files)} file(s){' + URL sources' if has_urls else ''}")

    # 2. Classify and extract ──────────────────────────────────────────
    # Each entry is (text, source_filename)
    sourced_text_parts = []
    image_items = []             # (label, b64, mime) for Gemini vision
    image_sources = {}           # label → source filename

    for filename in all_files:
        file_path = os.path.join(PDF_DIR, filename)
        ext = os.path.splitext(filename)[1].lower()
        base_name = os.path.splitext(filename)[0]
        # Clean display name (strip chatbot_id prefix)
        display_name = filename[len(chatbot_id) + 1:] if filename.startswith(chatbot_id + "_") else filename

        try:
            if ext == '.pdf':
                print(f"  📄 PDF: {filename}")
                text, scanned_pages = _extract_text_from_pdf(file_path)
                if text.strip():
                    sourced_text_parts.append((text, display_name))
                    print(f"     ✓ Extracted text from {filename} ({len(text)} chars)")
                for label, b64 in scanned_pages:
                    full_label = f"{base_name}_{label}"
                    image_items.append((full_label, b64, "image/png"))
                    image_sources[full_label] = display_name
                if scanned_pages:
                    print(f"     ⚡ {len(scanned_pages)} scanned page(s) queued for vision")

            elif ext == '.docx':
                print(f"  📝 DOCX: {filename}")
                text = _extract_text_from_docx(file_path)
                if text.strip():
                    sourced_text_parts.append((text, display_name))
                    print(f"     ✓ Extracted text ({len(text)} chars)")
                else:
                    print(f"     ⚠️ No text found in {filename}")

            elif ext == '.doc':
                print(f"  📝 DOC (legacy): {filename}")
                text = _convert_doc_to_text(file_path)
                if text:
                    sourced_text_parts.append((text, display_name))
                    print(f"     ✓ Extracted text ({len(text)} chars)")
                else:
                    print(f"     ⚠️ Could not extract text from {filename}")

            elif ext == '.txt':
                print(f"  📃 TXT: {filename}")
                text = _extract_text_from_txt(file_path)
                if text.strip():
                    sourced_text_parts.append((text, display_name))
                    print(f"     ✓ Read text ({len(text)} chars)")

            elif ext in IMAGE_EXTENSIONS:
                print(f"  🖼️  Image: {filename}")
                b64, mime = _image_to_base64(file_path)
                image_items.append((base_name, b64, mime))
                image_sources[base_name] = display_name

            else:
                print(f"  ⚠️ Skipping unsupported file: {filename}")

        except Exception as e:
            logger.exception("Error processing %s", filename)
            print(f"  ❌ Error processing {filename}: {e}")

    # 3. Gemini vision for images + scanned pages ──────────────────────
    if image_items:
        print(f"\n  🔍 Running Gemini vision on {len(image_items)} image(s) …")
        vision_text = _caption_images_concurrent(image_items, max_workers=3)
        if vision_text.strip():
            # Attribute vision text to the first image source (best-effort)
            first_source = image_sources.get(image_items[0][0], "images")
            sourced_text_parts.append((vision_text, first_source))
        print(f"  ✓ Vision processing complete")
        # Track vision API usage for cost monitoring
        try:
            from user_querySafe.models import VisionAPIUsage
            VisionAPIUsage.objects.create(
                chatbot_id=chatbot_id,
                call_count=len(image_items),
                call_type='training'
            )
        except Exception:
            pass  # Non-critical — don't fail pipeline over tracking

    # 3b. URL content ───────────────────────────────────────────────────
    try:
        from user_querySafe.models import ChatbotURL
        url_records = ChatbotURL.objects.filter(chatbot__chatbot_id=chatbot_id)
        if url_records.exists():
            print(f"\n  🌐 Processing URL sources …")
            from user_querySafe.chatbot.url_scraper import crawl_urls, parse_sitemap

            all_page_urls = []
            for record in url_records:
                if record.is_sitemap:
                    discovered_urls, err = parse_sitemap(record.url)
                    if err:
                        record.status = 'error'
                        record.error_message = err
                        record.save()
                        logger.warning("Sitemap error for %s: %s", record.url, err)
                    else:
                        record.page_count = len(discovered_urls)
                        record.status = 'crawled'
                        record.save()
                        all_page_urls.extend(discovered_urls)
                        print(f"     Sitemap: {len(discovered_urls)} pages from {record.url[:60]}")
                else:
                    all_page_urls.append(record.url)

            # Deduplicate while preserving order
            seen = set()
            unique_urls = []
            for u in all_page_urls:
                if u not in seen:
                    seen.add(u)
                    unique_urls.append(u)

            if unique_urls:
                crawl_results = crawl_urls(unique_urls, max_pages=50)
                url_text_count = 0
                for result in crawl_results:
                    if result['content']:
                        sourced_text_parts.append((result['content'], result['url']))
                        url_text_count += 1
                    elif result['error']:
                        logger.warning("URL crawl error for %s: %s", result['url'], result['error'])

                # Update non-sitemap URL record statuses
                for record in url_records.filter(is_sitemap=False):
                    matching = [r for r in crawl_results if r['url'] == record.url]
                    if matching:
                        if matching[0]['error']:
                            record.status = 'error'
                            record.error_message = matching[0]['error']
                        else:
                            record.status = 'crawled'
                        record.save()

                print(f"  ✓ Extracted content from {url_text_count}/{len(unique_urls)} URL(s)")
    except ImportError:
        logger.debug("ChatbotURL model or url_scraper not available, skipping URL processing")
    except Exception as e:
        logger.warning("URL processing error: %s", e)

    # 4. Combine, chunk, and track source per chunk ────────────────────
    if not sourced_text_parts:
        print(f"  ❌ No text extracted from any file for chatbot {chatbot_id}")
        from user_querySafe.models import Chatbot
        Chatbot.objects.filter(chatbot_id=chatbot_id).update(status="error")
        return

    # Chunk each source separately so we can tag chunks with their origin
    chunk_records = []  # [{"content": str, "source": str}, ...]
    combined_text_parts = []
    for text, source in sourced_text_parts:
        combined_text_parts.append(text)
        source_chunks = _chunk_text(text)
        for c in source_chunks:
            chunk_records.append({"content": c, "source": source})

    combined_text = "\n\n".join(combined_text_parts)
    print(f"\n  Total extracted text: {len(combined_text)} chars")
    print(f"  Chunked into {len(chunk_records)} segments")

    # Save combined text & chunks to disk (useful for debugging)
    text_path = os.path.join(TEXT_DIR, f"{chatbot_id}.txt")
    with open(text_path, "w", encoding="utf-8") as f:
        f.write(combined_text)

    chunk_path = os.path.join(CHUNK_DIR, f"{chatbot_id}-chunks.txt")
    with open(chunk_path, "w", encoding="utf-8") as f:
        for idx, rec in enumerate(chunk_records, 1):
            f.write(f"--- Chunk {idx} [{rec['source']}] ---\n{rec['content']}\n\n")

    # 5. Embed & index ─────────────────────────────────────────────────
    success = _embed_and_index(chatbot_id, chunk_records)

    # 6. Update chatbot status ─────────────────────────────────────────
    from user_querySafe.models import Chatbot
    chatbot_obj = Chatbot.objects.get(chatbot_id=chatbot_id)

    if success:
        chatbot_obj.status = "trained"
        chatbot_obj.dataset_name = f"{chatbot_id}-index.index"
        chatbot_obj.last_trained_at = timezone.now()
        chatbot_obj.save()
        elapsed = time.time() - start_time
        print(f"\n🎉 Pipeline completed for {chatbot_id} in {elapsed:.1f}s")
        print(f"   • Status: trained")
        print(f"   • Dataset: {chatbot_obj.dataset_name}")
    else:
        chatbot_obj.status = "error"
        chatbot_obj.save()
        print(f"\n❌ Pipeline failed — no embeddings produced for {chatbot_id}")


# =====================================================================
# BACKGROUND RUNNER (with simple lock to prevent duplicate runs)
# =====================================================================
_pipeline_locks = {}

def run_pipeline_background(chatbot_id):
    from threading import Thread, Lock

    lock = _pipeline_locks.setdefault(chatbot_id, Lock())
    if lock.locked():
        print(f"Pipeline already running for chatbot {chatbot_id}.")
        return

    def pipeline_wrapper():
        with lock:
            try:
                process_pipeline(chatbot_id)
            except Exception as e:
                logger.exception("Pipeline failed for chatbot %s", chatbot_id)
                try:
                    from user_querySafe.models import Chatbot
                    Chatbot.objects.filter(chatbot_id=chatbot_id).update(status='failed')
                except Exception:
                    logger.exception("Could not update chatbot %s status", chatbot_id)

    t = Thread(target=pipeline_wrapper, daemon=False)
    t.start()
